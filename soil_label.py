import csv
from docx import Document
from docx.shared import Inches, Pt

# Read the CSV file
def read_csv_file(file_path):
    with open(file_path, newline='') as csvfile:
        reader = csv.DictReader(csvfile)
        data = [row for row in reader]
    return data

# Add data to the table cell
def add_data(cell, data):
    # Add and format the Name
    name = cell.paragraphs[0].add_run(data['Name'])
    name.font.bold = True
    name.font.size = Pt(12)

    # Add and format the Field Name
    field_name = cell.paragraphs[0].add_run(f"\n{data['Field Name']}")
    field_name.font.size = Pt(12)

    # Add and format the Acres
    acres = cell.paragraphs[0].add_run(f"\nAcres: {data['Acres']}")
    acres.font.italic = True
    acres.font.size = Pt(12)

    # Add and format the New Crop
    new_crop = cell.paragraphs[0].add_run(f"\nNew Crop: {data['New Crop']}")
    new_crop.font.size = Pt(12)

# Create the Word document with the data
def create_word_document(data, output_path):
    doc = Document()

    # Set up page size and margins (assuming Avery 58160)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.1875)
    section.right_margin = Inches(0.1875)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    # Create a table for the labels
    table = doc.add_table(rows=10, cols=3)
    
    # Set up column widths and row heights
    for row in table.rows:
        row.height = Inches(1)
    for col in table.columns:
        col.width = Inches(2.625)

    # Add data to the table
    i = 0
    for row in table.rows:
        for cell in row.cells:
            if i < len(data):
                add_data(cell, data[i])
                i += 1

    # Save the document
    doc.save(output_path)

if __name__ == '__main__':
    csv_file_path = r'csv file here'  # Replace with the path to your CSV file
    output_path = r'output path'# Replace with the desired output path
    data = read_csv_file(csv_file_path)
    create_word_document(data, output_path)
